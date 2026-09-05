# -*- coding: utf-8 -*-
"""
build_docx.py — 论文.md → 复刻 C050 国一排版的 .docx（v3，2026-08-27 二轮迭代）
========================================================================
把 Markdown 论文转成对标 2025 C 题 C050 国一实测排版的 Word：
  - 页面 A4 / 边距 2.70(左右) 2.54(上下) cm；页脚居中页码（PAGE 域）
  - 正文 宋体/Times New Roman 小四 12pt 黑字、单倍行距、段前段后 0、首行缩进 2 字符
    （C050 实测：相邻基线差 15~16pt 即 Word 单倍行距；段落间无空行空隙）
  - 首行 `#` 为论文题名：黑体三号 16pt 居中；"摘  要"标题：黑体 16pt 居中；
    "关键词："行黑体引领；一级标题黑体四号 14pt、二级/三级黑体小四 12pt
  - 重点加粗范式（C050）：**bold** 正常转为加粗 run——摘要方法名/关键数字适度加粗、
    正文关键结论短语克制加粗（密度指引见 references/17 v3）
  - 无封面信息表（C050 首页直接题名+摘要+关键词）
  - 三线表（顶/底粗、表头下细、无竖线），表内文字宋体五号 10.5pt；
    图题在下、表题在上，图表题黑体五号 10.5pt
  - 行间公式 $$...$$ (n)：LaTeX → MathML → OMML 真公式对象（Word 公式编辑器原生），
    制表位实现"公式居中 + 编号右对齐"；转换管线 latex2mathml + Office MML2OMML.XSL，
    XSL 缺失时尝试 mathml2omml 纯 Python 包，两者皆不可用才降级为平文并打印警告
  - ``` 围栏代码块：Consolas 等宽 9pt 单倍行距（附录代码版式）

用法：
  python build_docx.py 论文.md --out 论文.docx [--template 国赛模板.docx]
约定（写 md 时）：
  # 一级 / ## 二级 / ### 三级（首个 # 记为论文题名）
  ![图 1 销售量季节性波动](figures/fig1.svg)   ← alt 即图题（在下）
  表 1 各品类描述统计量                 ← 表前一行以"表"开头即表题（在上）
  | 符号 | 说明 | 单位 |                  ← Markdown 表 → 三线表
  $$...$$ (1)                             ← 行间公式，(n) 为右侧编号
  **关键结论**                            ← 克制加粗（摘要≥300字符、正文每页1~4处）
  [1] 作者.题名[类型].刊名,年,卷(期):页码.   ← 参考文献
依赖：python-docx；公式对象化另需 latex2mathml+lxml（推荐）或 mathml2omml
"""
import sys, re, argparse, os, glob, copy
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BLACK = RGBColor(0, 0, 0)
HEI = "黑体"; SONG = "宋体"; TNR = "Times New Roman"
MONO = "Consolas"

# ---------------- 行间公式：LaTeX → MathML → OMML ----------------
_TEXT_WIDTH_CM = 15.6   # A4 21cm − 左右边距 2.7cm×2

_omml_xslt = None
_omml_mode = None       # 'xslt' | 'purepy' | None

def _find_mml2omml_xsl():
    pats = []
    for base in (r"C:\Program Files\Microsoft Office", r"C:\Program Files (x86)\Microsoft Office"):
        pats.append(os.path.join(base, "root", "Office*", "MML2OMML.XSL"))
        pats.append(os.path.join(base, "Office*", "MML2OMML.XSL"))
    for pat in pats:
        for hit in glob.glob(pat):
            if os.path.exists(hit):
                return hit
    return None

def _get_omml_transform():
    """懒加载转换器：优先 Office MML2OMML.XSL（lxml XSLT），否则 mathml2omml 纯 Python。"""
    global _omml_xslt, _omml_mode
    if _omml_mode is not None:
        return _omml_mode
    xsl = _find_mml2omml_xsl()
    if xsl:
        from lxml import etree
        _omml_xslt = etree.XSLT(etree.parse(xsl))
        _omml_mode = "xslt"
        return _omml_mode
    try:
        import mathml2omml  # noqa: F401
        _omml_mode = "purepy"
        return _omml_mode
    except ImportError:
        _omml_mode = None
        return None

def latex_to_omml_element(latex):
    """LaTeX → m:oMath 的 lxml 元素；失败抛异常，由调用方降级。"""
    import latex2mathml.converter as l2m
    mathml = l2m.convert(latex)
    mode = _get_omml_transform()
    if mode == "xslt":
        from lxml import etree
        result = _omml_xslt(etree.fromstring(mathml))
        root = result.getroot()
        if not root.tag.endswith('}oMath'):
            found = root.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath')
            if not found:
                raise RuntimeError("XSLT 输出不含 m:oMath")
            root = found[0]
        return copy.deepcopy(root)
    elif mode == "purepy":
        import mathml2omml
        from lxml import etree
        omml_str = mathml2omml.convert(mathml)
        return etree.fromstring(omml_str)
    raise RuntimeError("无可用的 MathML→OMML 转换器（缺 MML2OMML.XSL 且未安装 mathml2omml）")

FORMULA_RE = re.compile(r'^\$\$(.+)\$\$\s*\((\d+)\)\s*$')

def _no_indent(p):
    """段落级清零首行缩进（覆盖 Normal 的 2 字符缩进；同时清 firstLineChars 防继承合并）"""
    ppr = p._p.get_or_add_pPr()
    ind = ppr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind'); ppr.append(ind)
    ind.set(qn('w:firstLine'), '0')
    ind.set(qn('w:firstLineChars'), '0')

def _set_spacing(ppr, before='0', after='0', line='240'):
    sp = ppr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing'); ppr.append(sp)
    sp.set(qn('w:before'), before); sp.set(qn('w:after'), after)
    sp.set(qn('w:line'), line); sp.set(qn('w:lineRule'), 'auto')

def add_formula_paragraph(doc, latex, number):
    """公式段：居中制表位放公式对象 + 右对齐制表位放编号。返回是否对象化成功。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Pt(0)
    _no_indent(p)
    pf.tab_stops.add_tab_stop(Cm(_TEXT_WIDTH_CM / 2), WD_TAB_ALIGNMENT.CENTER)
    pf.tab_stops.add_tab_stop(Cm(_TEXT_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT)
    r_tab = p.add_run("\t"); set_run_font(r_tab, size=12)
    try:
        p._p.append(latex_to_omml_element(latex))
    except Exception as e:
        print(f"[公式降级] 式({number}) 转换失败({e})，以平文呈现")
        r = p.add_run(latex); set_run_font(r, size=12)
        r.font.italic = True
    r_tab2 = p.add_run("\t"); set_run_font(r_tab2, size=12)
    r_num = p.add_run(f"({number})"); set_run_font(r_num, size=12)
    return True

def ensure_code_style(doc):
    """附录代码样式：Consolas + 宋体 9pt、单倍行距、无首行缩进。"""
    try:
        st = doc.styles['代码']
    except KeyError:
        st = doc.styles.add_style('代码', WD_STYLE_TYPE.PARAGRAPH)
    st.base_style = doc.styles['Normal']
    st.font.name = MONO; st.font.size = Pt(9); st.font.color.rgb = BLACK
    rpr = st.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:eastAsia'), SONG); rf.set(qn('w:ascii'), MONO); rf.set(qn('w:hAnsi'), MONO)
    ppr = st.element.get_or_add_pPr()
    ind = ppr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind'); ppr.append(ind)
    ind.set(qn('w:firstLine'), '0'); ind.set(qn('w:firstLineChars'), '0')
    spacing = ppr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing'); ppr.append(spacing)
    spacing.set(qn('w:line'), '240'); spacing.set(qn('w:lineRule'), 'auto')
    return st

def add_code_line(doc, text):
    p = doc.add_paragraph()
    p.style = doc.styles['代码']
    r = p.add_run(text if text else "")
    set_run_font(r, cn=SONG, en=MONO, size=9)
    return p

def set_run_font(run, cn=SONG, en=TNR, size=10.5, bold=False, color=BLACK):
    run.font.name = en
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), cn)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)

def style_normal(doc):
    """Normal = 宋体/TNR 小四 12pt、单倍行距、段前段后 0、首行缩进 2 字符（C050 实测）"""
    st = doc.styles['Normal']
    st.font.name = TNR; st.font.size = Pt(12); st.font.color.rgb = BLACK
    rpr = st.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:eastAsia'), SONG)
    # 单倍行距 + 段前段后 0（C050：相邻基线差 15~16pt，段落间无空隙）
    ppr = st.element.get_or_add_pPr()
    _set_spacing(ppr, line='240')
    # 首行缩进 2 字符（12pt × 2 = 24pt = 480 twip）
    ind = ppr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind'); ppr.append(ind)
    ind.set(qn('w:firstLine'), '480')

def ensure_heading_style(doc, name, size, cn=HEI):
    """黑体标题样式：字号按级（H1=14 四号 / H2/H3=12 小四），单倍行距、无缩进"""
    try:
        st = doc.styles[name]
    except KeyError:
        st = doc.styles.add_style(name, 1)  # 1=PARAGRAPH
    st.font.name = TNR; st.font.size = Pt(size); st.font.bold = False; st.font.color.rgb = BLACK
    rpr = st.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:eastAsia'), cn)
    ppr = st.element.get_or_add_pPr()
    jc = ppr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc'); ppr.append(jc)
    jc.set(qn('w:val'), 'left')
    _set_spacing(ppr)   # 标题亦连续排版：段前段后 0、单倍行距
    ind = ppr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind'); ppr.append(ind)
    ind.set(qn('w:firstLine'), '0'); ind.set(qn('w:firstLineChars'), '0')

def ensure_caption_style(doc):
    try:
        st = doc.styles['图表标题']
    except KeyError:
        st = doc.styles.add_style('图表标题', 1)
    st.font.name = TNR; st.font.size = Pt(10.5); st.font.bold = True; st.font.color.rgb = BLACK
    rpr = st.element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:eastAsia'), HEI)
    ppr = st.element.get_or_add_pPr()
    jc = ppr.find(qn('w:jc'))
    if jc is None:
        jc = OxmlElement('w:jc'); ppr.append(jc)
    jc.set(qn('w:val'), 'center')
    _set_spacing(ppr)
    ind = ppr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind'); ppr.append(ind)
    ind.set(qn('w:firstLine'), '0'); ind.set(qn('w:firstLineChars'), '0')

def add_footer_pagenum(doc):
    """页脚居中页码（PAGE 域，宋体/TNR 五号）——C050 页脚居中页码"""
    sec = doc.sections[0]
    footer = sec.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_indent(p)
    _set_spacing(p._p.get_or_add_pPr())
    def _fld(kind=None, instr=None):
        r = p.add_run()
        if kind:
            el = OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'), kind)
        else:
            el = OxmlElement('w:instrText'); el.set(qn('xml:space'), 'preserve'); el.text = instr
        r._element.append(el)
        set_run_font(r, cn=SONG, en=TNR, size=10.5)
    _fld(kind='begin'); _fld(instr=' PAGE '); _fld(kind='end')

def set_three_line(table):
    """顶/底粗线、表头下细线、无竖线（17 §4）"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    def edge(tag, val, sz, color='000000'):
        e = OxmlElement(f'w:{tag}')
        e.set(qn('w:val'), val); e.set(qn('w:sz'), str(sz)); e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        borders.append(e)
    edge('top','single',12); edge('bottom','single',12)
    edge('left','none',0); edge('right','none',0)
    edge('insideV','none',0); edge('insideH','none',0)
    # 表头下分隔线放在 header row 的 trPr（单独细线）
    tblPr.append(borders)
    # header bottom thin line
    try:
        hdr = table.rows[0]._tr
        trPr = hdr.get_or_add_trPr()
        hb = OxmlElement('w:tblCellBorders')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'6'); bottom.set(qn('w:space'),'0'); bottom.set(qn('w:color'),'000000')
        hb.append(bottom); trPr.append(hb)
    except Exception:
        pass

def add_caption(doc, text, is_fig):
    p = doc.add_paragraph()
    p.style = doc.styles['图表标题']
    r = p.add_run(text); set_run_font(r, cn=HEI, en=TNR, size=10.5, bold=True)
    return p

def add_title_para(doc, text):
    """论文题名：黑体三号 16pt 居中（C050 首页范式）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_indent(p); _set_spacing(p._p.get_or_add_pPr())
    r = p.add_run(text); set_run_font(r, cn=HEI, en=TNR, size=16, bold=True)
    return p

def add_abstract_heading(doc, text='摘  要'):
    """摘要标题：黑体 16pt 居中"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_indent(p); _set_spacing(p._p.get_or_add_pPr())
    r = p.add_run(text); set_run_font(r, cn=HEI, en=TNR, size=16, bold=True)
    return p

def add_keywords_para(doc, text):
    """关键词行："关键词："黑体引领 + 关键词正文（宋体/TNR 12pt）"""
    p = doc.add_paragraph()
    _no_indent(p); _set_spacing(p._p.get_or_add_pPr())
    m = re.match(r'^(关键词\s*[：:])(.*)$', text)
    if m:
        r = p.add_run(m.group(1)); set_run_font(r, cn=HEI, en=TNR, size=12, bold=True)
        add_md_runs(p, m.group(2), size=12)
    else:
        add_md_runs(p, text, size=12)
    return p

def add_md_runs(p, text, size=12):
    """把含 **加粗** 的行内文本写入段落：`**x**` 转加粗 run，其余常规 run；
    行内反引号视为普通标记去除（C050 重点加粗范式）"""
    text = text.replace('`', '')
    for part in re.split(r'(\*\*.*?\*\*)', text):
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            r = p.add_run(part[2:-2]); set_run_font(r, bold=True, size=size)
        else:
            r = p.add_run(part); set_run_font(r, size=size)

IMG_RE = re.compile(r'!\[(.*?)\]\((.*?)\)')
TBL_RE = re.compile(r'^\s*\|.*\|\s*$')

def is_md_table(lines, i):
    if not TBL_RE.match(lines[i]): return False
    if i+1 < len(lines) and re.match(r'^\s*\|?[\s:|-]+\|?\s*$', lines[i+1]) and '-' in lines[i+1]:
        return True
    return False

def parse_table(rows):
    def split(line):
        line=line.strip().strip('|')
        return [c.strip() for c in line.split('|')]
    header = split(rows[0])
    data = [split(r) for r in rows[2:]]
    return header, data

def run():
    ap = argparse.ArgumentParser()
    ap.add_argument("md")
    ap.add_argument("--out", default="论文.docx")
    ap.add_argument("--template", default=None)
    args = ap.parse_args()

    if args.template:
        doc = Document(args.template)
        # 清空模板批注正文（保留样式定义）：删除所有段落/表格，重建
        # 简单起见：保留样式，清空 body 内容
        from docx.oxml import parse_xml
        body = doc.element.body
        for child in list(body):
            if child.tag == qn('w:sectPr'):
                continue
            body.remove(child)
    else:
        doc = Document()
        sec = doc.sections[0]
        sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
        sec.left_margin = Cm(2.7); sec.right_margin = Cm(2.7)
        sec.top_margin = Cm(2.54); sec.bottom_margin = Cm(2.54)

    style_normal(doc)
    ensure_heading_style(doc, 'Heading 1', 14)
    ensure_heading_style(doc, 'Heading 2', 12)
    ensure_heading_style(doc, 'Heading 3', 12)
    ensure_caption_style(doc)
    ensure_code_style(doc)
    add_footer_pagenum(doc)   # 页脚居中页码（C050 范式；v3 起不再生成封面信息表）

    lines = open(args.md, encoding='utf-8').read().splitlines()
    fig_n = tab_n = eq_n = 0
    title_done = False
    i = 0
    n = len(lines)
    in_fence = False
    while i < n:
        line = lines[i]
        s = line.strip()
        # ``` 围栏代码块：整块以等宽代码样式逐行排版（附录代码版式）
        if s.startswith('```'):
            in_fence = not in_fence
            i += 1; continue
        if in_fence:
            add_code_line(doc, line.rstrip())
            i += 1; continue
        if not s:
            i += 1; continue
        # 行间公式 $$...$$ (n)
        m = FORMULA_RE.match(s)
        if m:
            add_formula_paragraph(doc, m.group(1).strip(), m.group(2))
            eq_n += 1
            i += 1; continue
        # 标题：首个标题行为论文题名（黑体三号16pt居中）；"摘要"为黑体16pt居中
        m = re.match(r'^(#{1,3})\s+(.*)$', s)
        if m:
            level = len(m.group(1))
            txt = m.group(2).strip()
            if not title_done:
                title_done = True
                add_title_para(doc, txt)
                i += 1; continue
            if re.match(r'^摘\s*要$', txt):
                title_done = True
                add_abstract_heading(doc)
                i += 1; continue
            p = doc.add_paragraph()
            p.style = doc.styles[f'Heading {level}']
            r = p.add_run(txt); set_run_font(r, cn=HEI, en=TNR, size=14 if level==1 else 12, bold=False)
            i += 1; continue
        # 关键词行：黑体引领
        if re.match(r'^关键词\s*[：:]', s):
            add_keywords_para(doc, s)
            i += 1; continue
        # 图片
        m = IMG_RE.match(s)
        if m:
            cap, path = m.group(1), m.group(2)
            fig_n += 1
            # 图题文本：若 alt 无"图 N"，自动补
            cap_text = cap if cap.startswith("图") else f"图 {fig_n} {cap}"
            try:
                doc.add_picture(path, width=Cm(14))
                pic_p = doc.paragraphs[-1]
                pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _no_indent(pic_p); _set_spacing(pic_p._p.get_or_add_pPr())
            except Exception as e:
                doc.add_paragraph(f"[图丢失: {path} ({e})]")
            add_caption(doc, cap_text, is_fig=True)
            i += 1; continue
        # Markdown 表
        if is_md_table(lines, i):
            j = i
            while j < n and TBL_RE.match(lines[j]): j += 1
            header, data = parse_table(lines[i:j])
            tab_n += 1
            t = doc.add_table(rows=1, cols=len(header))
            def _fill_cell(c, text, bold=False):
                c.text=""
                cp = c.paragraphs[0]
                _no_indent(cp); _set_spacing(cp._p.get_or_add_pPr())
                add_md_runs(cp, text, size=10.5)
                if bold:
                    for r_ in cp.runs: r_.font.bold = True
            for c, h in zip(t.rows[0].cells, header):
                _fill_cell(c, h, bold=True)
            for row in data:
                for c, v in zip(t.add_row().cells, row):
                    _fill_cell(c, v)
            set_three_line(t)
            i = j; continue
        # 表题行（以"表"开头、单独成行）：表题在上（图表标题样式）
        if re.match(r'^表\s*\d+', s):
            add_caption(doc, s, is_fig=False)
            i += 1; continue
        # 参考文献行（12pt、无首行缩进）
        if re.match(r'^\[\d+\]', s):
            p = doc.add_paragraph()
            _no_indent(p); _set_spacing(p._p.get_or_add_pPr())
            r = p.add_run(s); set_run_font(r, size=12)
            i += 1; continue
        # 普通段落：**bold** → 加粗 run，`…` 去标记
        p = doc.add_paragraph()
        add_md_runs(p, s, size=12)
        i += 1
    # 文档元数据（对齐 CUMCMThesis 模板做法；作者置空——竞赛匿名要求，防文件属性泄露身份）
    cp = doc.core_properties
    cp.author = ""
    cp.last_modified_by = ""
    cp.subject = "全国大学生数学建模竞赛参赛论文"
    _t = next((l.lstrip('# ').strip() for l in lines if l.startswith('# ')), "")
    if _t:
        cp.title = _t
    doc.save(args.out)
    print(f"已生成 {args.out}（图 {fig_n} 张，表 {tab_n} 张，公式 {eq_n} 条）")

if __name__ == "__main__":
    run()
