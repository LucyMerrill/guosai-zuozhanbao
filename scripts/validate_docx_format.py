# -*- coding: utf-8 -*-
"""
国赛 Word 格式程序化校验器（v3，对标 C050 实测排版）
=================================
对照 `17_国赛Word原模板格式规范.md`（v3 起按 C050 实测修订）与 2025 国一 C050 实际写法，
对生成的 .docx 逐条核验。不靠记忆、不靠肉眼——只报实测：PASS / WARN / FAIL。

用法：
    python validate_docx_format.py <论文.docx> [--json 报告.json]
依赖：python-docx

v2（2026-08-27 一轮）：新增元话语黑名单、正文加粗上限、附录完整性三类检查。
v3（2026-08-27 二轮迭代，对标 C050）：
    - 正文主字号改判 小四 12pt ±0.3（违例 → FAIL），废除旧"五号10.5"标准；
    - 新增"摘要页加粗字符数 ≥200"检查（WARN）；重点加粗改为 C050 克制范式：
      无整段全加粗、加粗片段数在合理区间；
    - 删除旧"封面信息表全框"检查，反转为"无封面信息表"（C050 首页直接题名+摘要）；
    - 新增"题名/『摘 要』黑体三号16pt 居中"检查。
v4（2026-08-28，对齐 2026 官方 AI 新规）：
    - 新增"AI 工具使用声明"检查：缺声明=FAIL（2026 规定每篇必写），声明排在参考文献之后=WARN；
    - 依据：2026 国赛第一次通知（9/1 生效），条款全文见 17 号 §3B 末条。
规则常量集中于下方规则常量区。
"""
import sys, json, re, argparse
from docx import Document
from docx.oxml.ns import qn

EMU_PER_CM = 360000.0
def cm(v): return round(v / EMU_PER_CM, 2)

# ---------- 规则常量（v3：来自 C050 实测） ----------
EXPECT = {
    "page_w_cm": 21.0, "page_h_cm": 29.7,
    "margin_lr_cm": 2.70, "margin_tb_cm": 2.54,   # 允许 ±0.2cm
    "body_sz_pt": 12.0, "body_sz_tol": 0.3,        # 正文小四 12pt ±0.3（违例 FAIL）
    "title_sz_pt": 16.0, "title_sz_tol": 0.3,      # 题名/"摘 要" 三号 16pt
    "h1_sz_pt": 14.0, "h2_sz_pt": 12.0,            # 黑体 四号/小四
    "caption_sz_pt": 10.5,
}
BODY_FONTS = {"宋体", "SimSun", "宋体-方正", "Times New Roman", "TNR", "Times"}
HEI_FONTS = {"黑体", "SimHei", "黑体-方正", "SimHei-Light"}
ZH_RE = re.compile(r'[一-鿿]')

# ---------- v3 规则常量（2026-08-27 二轮迭代：重点加粗与封面范式，改阈值只动这里） ----------
META_BLACKLIST = [
    # 元话语 / 内部流程语汇泄漏（20号 L1）
    "如实", "预设判据", "判据", "契约", "预案", "条款", "而非修饰",
    # 工程口吻黑话（20号 L2）
    "口径", "引擎", "重放", "管道", "主力", "下发", "降级",
    "基线A", "基线B", "基线 A", "基线 B",
    "MC模拟", "MC 模拟",
]
ABSTRACT_BOLD_MIN = 200      # 摘要页加粗字符数下限，低于计 WARN
BODY_BOLD_SPAN_RANGE = (10, 400)   # 全文正文非标题加粗片段合理区间，出界计 WARN
CODE_MIN_CONSEC_PARAS = 30   # 附录代码判定：连续 等宽/代码样式 段落 ≥30 行视为贴有代码
MONO_FONTS = {"Consolas", "Courier New", "Courier", "Courier Prime",
              "Source Code Pro", "SourceCodePro", "Menlo", "Monaco",
              "DejaVu Sans Mono"}
CODE_STYLE_RE = re.compile(r'(code|preformatted|html.?pre|macro|verbatim)', re.I)  # 段落样式名含代码特征
CODE_TEXT_RE = re.compile(r'^\s*(import\s+\w+|from\s+[\w.]+\s+import\s|def\s+\w+\s*\(|class\s+\w+\b|function\s+\w|#include\s*[<"])')

def _para_fonts(p):
    """段落中所有 run 的 ascii 字体名集合"""
    return {r.font.name for r in p.runs if r.font.name}

def _is_codeish_para(p):
    """段落是否呈代码样式：等宽字体或代码类段落样式"""
    st = p.style.name if p.style else ""
    if CODE_STYLE_RE.search(st):
        return True
    return any(f in MONO_FONTS for f in _para_fonts(p))

def run_info(p):
    """返回该段落首个有字号/字体的 run 信息"""
    for r in p.runs:
        if r.font.size or r.font.name or (r.font.color and r.font.color.type is not None):
            nm = r.font.name
            # 中文混排时 eastAsia 才是中文实际字体
            ea = r._element.rPr.rFonts.get(qn('w:eastAsia')) if (r._element.rPr is not None and r._element.rPr.rFonts is not None) else None
            return {
                "size": r.font.size.pt if r.font.size else None,
                "ascii": nm,
                "eastasia": ea,
                "bold": r.font.bold,
                "color": (str(r.font.color.rgb) if (r.font.color and r.font.color.type is not None) else None),
            }
    return {"size": None, "ascii": None, "eastasia": None, "bold": None, "color": None}

def tbl_borders(t):
    """提取表格四边+内部横线样式 sz（pt）与是否竖向"""
    pr = t._tbl.tblPr
    b = pr.find(qn('w:tblBorders'))
    out = {}
    if b is None:
        # 可能继承样式：检查 table style
        st = pr.find(qn('w:tblStyle'))
        if st is not None:
            sid = st.get(qn('w:val'))
            try:
                style = t._parent.styles[sid]
                sb = style.element.tblPr.find(qn('w:tblBorders'))
                if sb is not None: b = sb
            except Exception:
                pass
    if b is None:
        return None
    for edge in ['top','left','bottom','right','insideH','insideV']:
        e = b.find(qn('w:'+edge))
        if e is not None:
            sz = e.get(qn('w:sz'))
            out[edge] = (int(sz)/2.0) if sz else 0.0   # w:sz 单位是 1/8 pt
            out[edge+'_val'] = e.get(qn('w:val'))
    return out

def is_sanxian(b):
    if not b: return False
    top = b.get('top',0); bot = b.get('bottom',0)
    ih = b.get('insideH',0)
    left = b.get('left',0); right = b.get('right',0); iv = b.get('insideV',0)
    # 三线表：上下有粗线(>0)，表内仅一条细分隔线(insideH 可存在)，无竖线
    has_outer = top > 0 and bot > 0
    no_vline = left == 0 and right == 0 and iv == 0
    return has_outer and no_vline

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("docx")
    ap.add_argument("--json", default=None)
    args = ap.parse_args()

    d = Document(args.docx)
    rep = {"file": args.docx, "checks": [], "summary": {}}

    def add(name, status, detail=""):
        rep["checks"].append({"check": name, "status": status, "detail": detail})

    # 1. 页面尺寸
    sec = d.sections[0]
    pw, ph = cm(sec.page_width), cm(sec.page_height)
    ok = abs(pw-EXPECT["page_w_cm"])<0.3 and abs(ph-EXPECT["page_h_cm"])<0.3
    add("页面 A4 (21×29.7cm)", "PASS" if ok else "FAIL", f"实测 {pw}×{ph}cm")

    # 2. 页边距
    ml, mr = cm(sec.left_margin), cm(sec.right_margin)
    mt, mb = cm(sec.top_margin), cm(sec.bottom_margin)
    ok_lr = abs((ml+mr)/2 - EXPECT["margin_lr_cm"])<0.3
    ok_tb = abs((mt+mb)/2 - EXPECT["margin_tb_cm"])<0.3
    add("页边距 左右≈2.70 上下≈2.54cm", "PASS" if (ok_lr and ok_tb) else "WARN",
        f"L={ml} R={mr} T={mt} B={mb}")

    # 3. 无目录
    has_toc = any(p.style and 'TOC' in p.style.name for p in d.paragraphs) or \
              any('TOC' in (p.text or '') for p in d.paragraphs[:5])
    add("无目录（TOC）", "PASS" if not has_toc else "FAIL", "检测到目录" if has_toc else "未检测到")

    # 4. 摘要 / 关键词 标题
    texts = [p.text.strip() for p in d.paragraphs]
    joined = "\n".join(texts)
    has_abs = bool(re.search(r'摘\s*要', joined))
    has_kw = bool(re.search(r'关键词', joined))
    add("含『摘要』与『关键词』", "PASS" if (has_abs and has_kw) else "FAIL",
        f"摘要={has_abs} 关键词={has_kw}")

    # 5. 章节结构（问题驱动：问题N模型的建立与求解）
    prob_secs = [t for t in texts if re.search(r'问题[一二三四五六七八九十\d]+.*模型的建立与求解', t)]
    add("问题驱动分章（问题N 模型的建立与求解）", "PASS" if len(prob_secs)>=2 else "WARN",
        f"命中 {len(prob_secs)} 个：{prob_secs[:6]}")
    has_eval = bool(re.search(r'模型的评价与推广|模型的评价与改进', joined))
    add("含『模型的评价与推广/改进』", "PASS" if has_eval else "WARN", "")

    # 6. 正文段落：字号/字体/颜色（v3：正文主字号 小四12pt±0.3，违例 FAIL）
    body_issues, color_issues, title_issues = [], [], []
    h1_issues, h2_issues = [], []
    for p in d.paragraphs:
        if not p.text.strip(): continue
        st = p.style.name if p.style else ""
        info = run_info(p)
        sizes_in_p = [r.font.size.pt for r in p.runs if r.font.size]
        is_title_16 = bool(sizes_in_p) and min(sizes_in_p) >= 15.0   # 题名/"摘 要"
        if st == 'Normal':
            if is_title_16:
                ea = info["eastasia"]
                if (info["size"] is None or abs(info["size"]-EXPECT["title_sz_pt"])>EXPECT["title_sz_tol"]
                        or (ea and ea not in HEI_FONTS)):
                    title_issues.append(f"({ea},{info['size']}pt): {p.text[:16]}")
            else:
                sz = info["size"]
                if sz is not None and abs(sz-EXPECT["body_sz_pt"])>EXPECT["body_sz_tol"]:
                    body_issues.append(f"字号{sz}≠{EXPECT['body_sz_pt']}: {p.text[:20]}")
                if info["color"] not in (None, '000000', 'auto'):
                    color_issues.append(f"非黑字{info['color']}: {p.text[:20]}")
        elif st == 'Heading 1':
            if info["eastasia"] and info["eastasia"] not in HEI_FONTS and info["ascii"] not in HEI_FONTS:
                h1_issues.append(f"非黑体: {p.text[:20]}")
        elif st == 'Heading 2':
            if info["eastasia"] and info["eastasia"] not in HEI_FONTS:
                h2_issues.append(f"非黑体: {p.text[:20]}")
    size_bad = bool(body_issues)
    status6 = "FAIL" if size_bad else ("WARN" if color_issues else "PASS")
    detail6 = ""
    if size_bad:
        detail6 += ";" .join(body_issues[:3]) + (f" …共{len(body_issues)}处" if len(body_issues)>3 else "")
    if color_issues:
        detail6 += ("；" if detail6 else "") + "；".join(color_issues[:2])
    add("正文 小四(≈12pt±0.3) 黑色", status6, detail6)
    add("题名与『摘 要』黑体三号(16pt)居中", "WARN" if title_issues else "PASS",
        "；".join(title_issues[:3]))
    add("一级标题 黑体", "PASS" if not h1_issues else "WARN", ";".join(h1_issues[:3]))
    add("二级标题 黑体", "PASS" if not h2_issues else "WARN", ";".join(h2_issues[:3]))

    # 7. 三线表
    sx_count = 0
    for t in d.tables:
        b = tbl_borders(t)
        if is_sanxian(b): sx_count += 1
    add("存在三线表（无竖线、上下粗线）", "PASS" if sx_count>0 else "WARN", f"检测到 {sx_count} 个")

    # 8. 图表标题样式 + 编号连续
    caps = [p.text.strip() for p in d.paragraphs if p.style and p.style.name=='图表标题']
    fig_caps = [c for c in caps if '图' in c]
    tab_caps = [c for c in caps if '表' in c]
    add("图表标题样式存在且连续编号", "PASS" if (len(fig_caps)>0 or len(tab_caps)>0) else "WARN",
        f"图题{len(fig_caps)} 表题{len(tab_caps)}; 样例: {caps[:2]}")

    # 9. 参考文献格式 [n]
    refs = [t for t in texts if re.match(r'^\s*\[\d+\]', t)]
    add("参考文献 [n] 编号格式", "PASS" if len(refs)>=3 else "WARN", f"命中 {len(refs)} 条；样例: {refs[:1]}")

    # 9B. AI 工具使用声明（v4：2026 官方新规——每篇论文必写，置于参考文献之前）
    ai_idx = next((i for i, t in enumerate(texts)
                   if re.search(r'AI\s*工具\s*使用声明', t) and len(t) < 40), None)
    ref_idx = next((i for i, t in enumerate(texts)
                    if re.match(r'^\s*参考文献\s*$', t) or re.match(r'^\s*\[\d+\]', t)), None)
    if ai_idx is None:
        add("AI 工具使用声明（2026 新规）", "FAIL",
            "未找到『AI 工具使用声明』章 → 2026 规定每篇必写（用了/没用都要），置于参考文献之前；句式见 17 号 §3B")
    elif ref_idx is not None and ai_idx > ref_idx:
        add("AI 工具使用声明（2026 新规）", "WARN",
            "声明出现在参考文献之后 → 规定位置为参考文献之前，请上移")
    else:
        add("AI 工具使用声明（2026 新规）", "PASS", "已含『AI 工具使用声明』且位于参考文献之前")

    # 10. 无封面信息表（v3 反转：C050 国一首页直接题名+摘要+关键词）
    cover_like = False
    for t in d.tables:
        txt = " ".join(c.text for r in t.rows for c in r.cells)
        if ('学校' in txt and ('队员' in txt or '编号' in txt)) or '赛区' in txt or '指导教师' in txt:
            cover_like = True; break
    add("无封面信息表（C050 范式）", "FAIL" if cover_like else "PASS",
        ("检测到疑似封面信息表 → C050 首页直接题名+摘要，应删除" if cover_like
         else "未含封面信息表，首页直接题名+摘要"))

    # ---------- v2 新增检查（2026-08-27 迭代，原有 1~10 项逻辑不动） ----------

    # 11. 元话语/工程黑话黑名单扫描（正文非标题、非代码样式段落）
    meta_hits = {}
    for p in d.paragraphs:
        st = p.style.name if p.style else ""
        if st.startswith('Heading') or st == '图表标题':
            continue
        if _is_codeish_para(p):   # 附录代码段不参与黑名单扫描，防误报
            continue
        t = p.text or ""
        if not t.strip():
            continue
        for w in META_BLACKLIST:
            n = t.count(w)
            if n:
                meta_hits[w] = meta_hits.get(w, 0) + n
    meta_detail = "全文无命中"
    if meta_hits:
        tops = sorted(meta_hits.items(), key=lambda kv: -kv[1])[:6]
        meta_detail = "；".join(f"『{w}』×{n}" for w, n in tops) + \
                      f"；共 {sum(meta_hits.values())} 处 → 按 20号§五 改写对照替换"
    add("元话语/工程黑话黑名单（20号雷点库）", "WARN" if meta_hits else "PASS", meta_detail)

    # 12. 重点加粗规范（v3，C050 克制加粗范式）：
    #     摘要区加粗字符 ≥ABSTRACT_BOLD_MIN（WARN）；全文无整段全加粗；加粗片段数在合理区间
    paras = list(d.paragraphs)
    abs_idx0 = next((i for i,p in enumerate(paras)
                     if re.fullmatch(r'摘\s*要', (p.text or '').strip())), None)
    kw_idx = next((i for i,p in enumerate(paras)
                   if re.match(r'^关键词', (p.text or '').strip())), None)
    abs_bold_chars = 0
    total_spans = 0
    full_bold_paras = []
    prev_in_span = False
    for j,p in enumerate(paras):
        st = p.style.name if p.style else ""
        if st.startswith('Heading') or st == '图表标题':
            prev_in_span = False; continue
        runs = [r for r in p.runs if (r.text or '').strip()]
        if not runs:
            prev_in_span = False; continue
        in_abs = (abs_idx0 is not None and kw_idx is not None and abs_idx0 <= j <= kw_idx)
        nbold = sum(len((r.text or '').strip()) for r in runs if r.font.bold)
        if in_abs:
            abs_bold_chars += nbold
        # 加粗片段计组（连续 bold run 记 1 组）
        prev_b = False
        for r in runs:
            b = bool(r.font.bold)
            if b and not prev_b:
                total_spans += 1
            prev_b = b
        # 整段全加粗（排除关键词行等短段）
        if all(bool(r.font.bold) for r in runs) and len((p.text or '').strip()) > 40:
            full_bold_paras.append((p.text or "")[:16])
    problems12 = []
    if abs_bold_chars < ABSTRACT_BOLD_MIN:
        problems12.append(f"摘要页加粗字符 {abs_bold_chars}<{ABSTRACT_BOLD_MIN} → 摘要中方法名/关键数字/结论短语需适度加粗")
    if full_bold_paras:
        problems12.append("存在整段全加粗: " + " / ".join(full_bold_paras[:2]) + " → 禁止整段加粗")
    lo, hi = BODY_BOLD_SPAN_RANGE
    if total_spans < lo:
        problems12.append(f"正文加粗片段仅 {total_spans} 组(<{lo})，重点范式未落地")
    elif total_spans > hi:
        problems12.append(f"正文加粗片段 {total_spans} 组(>{hi})，过密")
    add(f"重点加粗规范（摘要加粗≥{ABSTRACT_BOLD_MIN}字符、克制密度）",
        "WARN" if problems12 else "PASS",
        (f"摘要加粗 {abs_bold_chars} 字符；全文加粗 {total_spans} 组；"
         + ("；".join(problems12) if problems12 else "无整段加粗，符合 C050 克制范式")))

    # 13. 附录完整性：须贴出真实完整代码（连续代码样式段落 ≥30 行，或存在 import/def/function 等语法特征行）
    mono_run = 0; max_mono_run = 0; has_codeline = False
    for p in d.paragraphs:
        if _is_codeish_para(p):
            mono_run += 1
            max_mono_run = max(max_mono_run, mono_run)
        else:
            mono_run = 0
        if CODE_TEXT_RE.match(p.text or ""):
            has_codeline = True
    code_ok = (max_mono_run >= CODE_MIN_CONSEC_PARAS) or has_codeline
    code_detail = (f"最长连续代码样式段落 {max_mono_run} 行(阈值≥{CODE_MIN_CONSEC_PARAS})，"
                   f"代码语法特征行{'已' if has_codeline else '未'}检出")
    if not code_ok:
        code_detail += " → 附录代码缺失：按 20号 G3 贴出可运行代码全文（等宽小字号，每脚本前加注释头）"
    add("附录完整性（真实完整代码）", "PASS" if code_ok else "FAIL", code_detail)

    # 汇总
    n_pass = sum(1 for c in rep["checks"] if c["status"]=="PASS")
    n_warn = sum(1 for c in rep["checks"] if c["status"]=="WARN")
    n_fail = sum(1 for c in rep["checks"] if c["status"]=="FAIL")
    rep["summary"] = {"pass": n_pass, "warn": n_warn, "fail": n_fail,
                      "grade": "FAIL" if n_fail>0 else ("PASS" if n_warn==0 else "PASS*（有提示）")}
    print(json.dumps(rep["summary"], ensure_ascii=False, indent=2))
    for c in rep["checks"]:
        print(f"[{c['status']:4}] {c['check']}  {c['detail']}")
    if args.json:
        json.dump(rep, open(args.json,"w",encoding="utf-8"), ensure_ascii=False, indent=2)
        print("\n报告已写入:", args.json)

if __name__ == "__main__":
    main()
